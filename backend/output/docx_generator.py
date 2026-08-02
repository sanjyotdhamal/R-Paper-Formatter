import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), '..', 'outputs')
os.makedirs(OUTPUT_DIR, exist_ok=True)

# =====================================================
# IEEE FORMATTING CONFIGURATION
# =====================================================
IEEE_CONFIG = {
    "page": {
        "paper_size": "A4",
        "width": "210mm",
        "height": "297mm",
        "top_margin": "19mm",
        "bottom_margin": "25.4mm",
        "left_margin": "14.32mm",
        "right_margin": "14.32mm",
        "orientation": "portrait"
    },
    "columns": {
        "enabled": True,
        "count": 2,
        "spacing": "4.22mm",
        "equal_width": True
    },
    "default_font": {
        "name": "Times New Roman",
        "size": 10,
        "color": "#000000"
    },
    "title": {
        "font": "Times New Roman",
        "size": 24,
        "bold": False,
        "italic": False,
        "alignment": "center",
        "line_spacing": 1.0,
        "space_before": 0,
        "space_after": 12,
        "all_caps": False
    },
    "authors": {
        "font": "Times New Roman",
        "size": 11,
        "bold": False,
        "alignment": "center",
        "space_before": 0,
        "space_after": 0
    },
    "affiliation": {
        "font": "Times New Roman",
        "size": 10,
        "bold": False,
        "italic": False,
        "alignment": "center",
        "space_before": 0,
        "space_after": 12
    },
    "abstract": {
        "heading": {
            "text": "Abstract",
            "font": "Times New Roman",
            "size": 9,
            "bold": True,
            "italic": False,
            "prefix_dash": "—"
        },
        "body": {
            "font": "Times New Roman",
            "size": 9,
            "bold": False,
            "italic": False,
            "alignment": "justify",
            "line_spacing": 1.0,
            "space_before": 0,
            "space_after": 8
        },
        "max_words": 250
    },
    "keywords": {
        "heading": {
            "text": "Keywords",
            "font": "Times New Roman",
            "size": 9,
            "bold": True,
            "italic": False,
            "prefix_dash": "—"
        },
        "body": {
            "font": "Times New Roman",
            "size": 9,
            "italic": False,
            "alignment": "justify",
            "separator": ", "
        },
        "space_after": 10
    },
    "body": {
        "font": "Times New Roman",
        "size": 10,
        "bold": False,
        "italic": False,
        "alignment": "justify",
        "line_spacing": 1.0,
        "first_line_indent": "0mm",
        "space_before": 0,
        "space_after": 0
    },
    "section_heading": {
        "font": "Times New Roman",
        "size": 10,
        "bold": True,
        "italic": False,
        "alignment": "center",
        "all_caps": True,
        "numbering_style": "roman_upper",
        "format": "I.",
        "space_before": 12,
        "space_after": 6
    },
    "subsection_heading": {
        "font": "Times New Roman",
        "size": 10,
        "bold": True,
        "italic": False,
        "alignment": "left",
        "numbering_style": "alphabet_upper",
        "format": "A.",
        "space_before": 6,
        "space_after": 3
    },
    "references": {
        "heading": {
            "text": "REFERENCES",
            "font": "Times New Roman",
            "size": 10,
            "bold": True,
            "alignment": "center",
            "all_caps": True
        },
        "body": {
            "font": "Times New Roman",
            "size": 8,
            "alignment": "justify",
            "line_spacing": 1.0,
            "hanging_indent": "3.7mm",
            "space_after": 2
        },
        "citation_style": "IEEE",
        "in_text_citation": "[1]",
        "numbering": "numeric"
    },
    "paragraphs": {
        "widow_control": True,
        "keep_with_next": False
    },
    "missing_content_handling": {
        "skip_if_empty": True,
        "placeholder_text": ""
    }
}


def mm_to_cm(mm_value):
    """Convert millimeters to centimeters"""
    return float(mm_value.replace("mm", "")) / 10


def set_font(run, font_name, font_size, bold=False, italic=False, color=None):
    """Apply font formatting to a run"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def set_paragraph_spacing(paragraph, space_before=0, space_after=0, line_spacing=1.0, alignment=None):
    """Apply spacing and alignment to a paragraph"""
    paragraph.paragraph_format.space_before = Pt(space_before)
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.line_spacing = line_spacing
    if alignment:
        paragraph.alignment = alignment


def get_alignment(alignment_str):
    """Convert alignment string to Word alignment constant"""
    alignment_map = {
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY
    }
    return alignment_map.get(alignment_str, WD_ALIGN_PARAGRAPH.LEFT)


def generate_ieee(content: dict, job_id: str) -> str:
    """Generate IEEE formatted document with proper formatting rules"""
    doc = Document()
    config = IEEE_CONFIG

    # ===== PAGE SETUP =====
    section = doc.sections[0]
    section.page_width = Cm(mm_to_cm(config["page"]["width"]))
    section.page_height = Cm(mm_to_cm(config["page"]["height"]))
    section.left_margin = Cm(mm_to_cm(config["page"]["left_margin"]))
    section.right_margin = Cm(mm_to_cm(config["page"]["right_margin"]))
    section.top_margin = Cm(mm_to_cm(config["page"]["top_margin"]))
    section.bottom_margin = Cm(mm_to_cm(config["page"]["bottom_margin"]))

    # ===== SETUP 2-COLUMN LAYOUT =====
    if config["columns"]["enabled"]:
        try:
            sectPr = section._sectPr
            cols = sectPr.xpath('./w:cols')[0] if sectPr.xpath('./w:cols') else OxmlElement('w:cols')
            cols.set(qn('w:num'), str(config["columns"]["count"]))
            cols.set(qn('w:sep'), '1')  # Add separator line
            if not sectPr.xpath('./w:cols'):
                sectPr.append(cols)
        except:
            pass  # If column setup fails, continue with single column

    # ===== TITLE =====
    title_cfg = config["title"]
    title_para = doc.add_paragraph()
    title_para.alignment = get_alignment(title_cfg["alignment"])
    title_run = title_para.add_run(content.get('title', 'Untitled Paper'))
    set_font(title_run, title_cfg["font"], title_cfg["size"], 
            bold=title_cfg["bold"], italic=title_cfg["italic"])
    set_paragraph_spacing(title_para, 
                        space_before=title_cfg["space_before"],
                        space_after=title_cfg["space_after"],
                        line_spacing=title_cfg["line_spacing"])

    # ===== AUTHORS =====
    if content.get('authors'):
        authors_cfg = config["authors"]
        authors_para = doc.add_paragraph()
        authors_para.alignment = get_alignment(authors_cfg["alignment"])
        authors = content.get('authors', [])
        authors_text = ', '.join(authors) if isinstance(authors, list) else authors
        authors_run = authors_para.add_run(authors_text)
        set_font(authors_run, authors_cfg["font"], authors_cfg["size"], 
                bold=authors_cfg["bold"])
        set_paragraph_spacing(authors_para,
                            space_before=authors_cfg["space_before"],
                            space_after=authors_cfg["space_after"])

    # ===== AFFILIATION =====
    if content.get('affiliation'):
        aff_cfg = config["affiliation"]
        aff_para = doc.add_paragraph()
        aff_para.alignment = get_alignment(aff_cfg["alignment"])
        aff_text = content.get('affiliation', '')
        if content.get('email'):
            aff_text += f"\n{content.get('email')}"
        aff_run = aff_para.add_run(aff_text)
        set_font(aff_run, aff_cfg["font"], aff_cfg["size"],
                bold=aff_cfg["bold"], italic=aff_cfg["italic"])
        set_paragraph_spacing(aff_para,
                            space_before=aff_cfg["space_before"],
                            space_after=aff_cfg["space_after"])

    # ===== ABSTRACT =====
    if content.get('abstract'):
        abs_cfg = config["abstract"]
        
        # Abstract heading with prefix dash
        abs_heading_para = doc.add_paragraph()
        abs_heading_run = abs_heading_para.add_run(abs_cfg["heading"]["prefix_dash"])
        set_font(abs_heading_run, abs_cfg["heading"]["font"], abs_cfg["heading"]["size"],
                bold=abs_cfg["heading"]["bold"], italic=abs_cfg["heading"]["italic"])
        
        abs_space_run = abs_heading_para.add_run(" ")
        set_font(abs_space_run, abs_cfg["heading"]["font"], abs_cfg["heading"]["size"])
        
        abs_head_text_run = abs_heading_para.add_run(abs_cfg["heading"]["text"])
        set_font(abs_head_text_run, abs_cfg["heading"]["font"], abs_cfg["heading"]["size"],
                bold=abs_cfg["heading"]["bold"], italic=abs_cfg["heading"]["italic"])
        
        # Abstract body on new line
        abstract_text = content.get('abstract', '')
        abs_text_run = abs_heading_para.add_run(f"\n{abstract_text}")
        set_font(abs_text_run, abs_cfg["body"]["font"], abs_cfg["body"]["size"],
                bold=abs_cfg["body"]["bold"], italic=abs_cfg["body"]["italic"])
        
        abs_heading_para.alignment = get_alignment(abs_cfg["body"]["alignment"])
        set_paragraph_spacing(abs_heading_para,
                            space_before=abs_cfg["body"]["space_before"],
                            space_after=abs_cfg["body"]["space_after"],
                            line_spacing=abs_cfg["body"]["line_spacing"])

    # ===== KEYWORDS =====
    if content.get('keywords'):
        kw_cfg = config["keywords"]
        kw_para = doc.add_paragraph()
        
        # Keywords heading
        kw_head_run = kw_para.add_run(kw_cfg["heading"]["prefix_dash"])
        set_font(kw_head_run, kw_cfg["heading"]["font"], kw_cfg["heading"]["size"],
                bold=kw_cfg["heading"]["bold"])
        
        kw_space_run = kw_para.add_run(" ")
        set_font(kw_space_run, kw_cfg["heading"]["font"], kw_cfg["heading"]["size"])
        
        kw_text_run = kw_para.add_run(kw_cfg["heading"]["text"])
        set_font(kw_text_run, kw_cfg["heading"]["font"], kw_cfg["heading"]["size"],
                bold=kw_cfg["heading"]["bold"])
        
        # Keywords body on new line
        keywords = content.get('keywords', [])
        keywords_text = kw_cfg["body"]["separator"].join(keywords) if isinstance(keywords, list) else keywords
        kw_body_run = kw_para.add_run(f"\n{keywords_text}")
        set_font(kw_body_run, kw_cfg["body"]["font"], kw_cfg["body"]["size"])
        
        kw_para.alignment = get_alignment(kw_cfg["body"]["alignment"])
        set_paragraph_spacing(kw_para, space_after=kw_cfg["space_after"])

    # ===== BODY SECTIONS =====
    section_cfg = config["section_heading"]
    body_cfg = config["body"]
    
    sections_data = [
        ('I. INTRODUCTION', 'introduction'),
        ('II. METHODOLOGY', 'methodology'),
        ('III. RESULTS', 'results'),
        ('IV. CONCLUSION', 'conclusion'),
    ]

    for heading_text, key in sections_data:
        section_content = content.get(key, '').strip()
        
        # Always add section heading if we have content
        if section_content:
            # Section heading
            heading_para = doc.add_paragraph()
            heading_para.alignment = get_alignment(section_cfg["alignment"])
            heading_run = heading_para.add_run(heading_text)
            set_font(heading_run, section_cfg["font"], section_cfg["size"],
                    bold=section_cfg["bold"], italic=section_cfg["italic"])
            set_paragraph_spacing(heading_para,
                                space_before=section_cfg["space_before"],
                                space_after=section_cfg["space_after"])

            # Section content
            content_para = doc.add_paragraph()
            content_para.alignment = get_alignment(body_cfg["alignment"])
            content_run = content_para.add_run(section_content)
            set_font(content_run, body_cfg["font"], body_cfg["size"],
                    bold=body_cfg["bold"], italic=body_cfg["italic"])
            set_paragraph_spacing(content_para,
                                space_before=body_cfg["space_before"],
                                space_after=body_cfg["space_after"],
                                line_spacing=body_cfg["line_spacing"])
            
            # Widow control
            if config["paragraphs"]["widow_control"]:
                content_para.paragraph_format.widow_control = True

    # ===== REFERENCES =====
    references = content.get('references', [])
    if references and len(references) > 0:
        ref_cfg = config["references"]
        
        # References heading
        ref_heading = doc.add_paragraph()
        ref_heading.alignment = get_alignment(ref_cfg["heading"]["alignment"])
        ref_run = ref_heading.add_run(ref_cfg["heading"]["text"])
        set_font(ref_run, ref_cfg["heading"]["font"], ref_cfg["heading"]["size"],
                bold=ref_cfg["heading"]["bold"])
        set_paragraph_spacing(ref_heading, space_before=12)

        # References body
        if isinstance(references, list):
            for i, ref in enumerate(references):
                if ref.strip():  # Only add non-empty references
                    ref_para = doc.add_paragraph()
                    ref_para.alignment = get_alignment(ref_cfg["body"]["alignment"])
                    ref_text = f"[{i+1}] {ref}"
                    ref_body_run = ref_para.add_run(ref_text)
                    set_font(ref_body_run, ref_cfg["body"]["font"], ref_cfg["body"]["size"])
                    set_paragraph_spacing(ref_para, 
                                        space_after=ref_cfg["body"]["space_after"],
                                        line_spacing=ref_cfg["body"]["line_spacing"])
                    
                    # Set hanging indent
                    ref_para.paragraph_format.left_indent = Cm(mm_to_cm(ref_cfg["body"]["hanging_indent"]))
                    ref_para.paragraph_format.first_line_indent = Cm(-mm_to_cm(ref_cfg["body"]["hanging_indent"]))
        else:
            ref_para = doc.add_paragraph()
            ref_para.alignment = get_alignment(ref_cfg["body"]["alignment"])
            ref_body_run = ref_para.add_run(references)
            set_font(ref_body_run, ref_cfg["body"]["font"], ref_cfg["body"]["size"])
            set_paragraph_spacing(ref_para,
                                space_after=ref_cfg["body"]["space_after"],
                                line_spacing=ref_cfg["body"]["line_spacing"])

    output_path = os.path.join(OUTPUT_DIR, f"{job_id}_ieee.docx")
    doc.save(output_path)
    return output_path


def generate_springer(content: dict, job_id: str) -> str:
    """Generate Springer formatted document"""
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(15.5)
    section.page_height = Cm(23.5)
    section.left_margin = Cm(4.3)
    section.right_margin = Cm(3.3)
    section.top_margin = Cm(6.3)
    section.bottom_margin = Cm(6.3)

    # Title
    if content.get('title'):
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        title_run = title_para.add_run(content.get('title', ''))
        set_font(title_run, 'Times New Roman', 14, bold=True)
        title_para.space_after = Pt(16)

    # Authors
    if content.get('authors'):
        authors_para = doc.add_paragraph()
        authors = content.get('authors', [])
        authors_text = ', '.join(authors) if isinstance(authors, list) else authors
        authors_run = authors_para.add_run(authors_text)
        set_font(authors_run, 'Times New Roman', 10)
        authors_para.space_after = Pt(16)

    # Abstract
    if content.get('abstract'):
        abstract_para = doc.add_paragraph()
        abstract_run = abstract_para.add_run('Abstract. ')
        set_font(abstract_run, 'Times New Roman', 10, bold=True)
        abstract_text_run = abstract_para.add_run(content.get('abstract', ''))
        set_font(abstract_text_run, 'Times New Roman', 10)
        abstract_para.space_after = Pt(8)

    # Keywords
    if content.get('keywords'):
        keywords_para = doc.add_paragraph()
        keywords = content.get('keywords', [])
        keywords_text = ', '.join(keywords) if isinstance(keywords, list) else keywords
        kw_bold = keywords_para.add_run('Keywords: ')
        set_font(kw_bold, 'Times New Roman', 10, bold=True)
        kw_run = keywords_para.add_run(keywords_text)
        set_font(kw_run, 'Times New Roman', 10)
        keywords_para.space_after = Pt(16)

    # Sections
    sections_data = [
        ('1 Introduction', 'introduction'),
        ('2 Methodology', 'methodology'),
        ('3 Results', 'results'),
        ('4 Conclusion', 'conclusion'),
    ]

    for heading_text, key in sections_data:
        if content.get(key) and content.get(key).strip():
            heading_para = doc.add_paragraph()
            heading_run = heading_para.add_run(heading_text)
            set_font(heading_run, 'Times New Roman', 12, bold=True)
            heading_para.space_before = Pt(8)
            heading_para.space_after = Pt(4)

            content_para = doc.add_paragraph()
            content_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            content_run = content_para.add_run(content.get(key, ''))
            set_font(content_run, 'Times New Roman', 10)
            content_para.space_after = Pt(6)

    # References
    if content.get('references'):
        ref_heading = doc.add_paragraph()
        ref_run = ref_heading.add_run('References')
        set_font(ref_run, 'Times New Roman', 12, bold=True)
        ref_heading.space_before = Pt(8)

        references = content.get('references', [])
        if isinstance(references, list):
            for i, ref in enumerate(references):
                ref_para = doc.add_paragraph()
                ref_run = ref_para.add_run(f"{i+1}. {ref}")
                set_font(ref_run, 'Times New Roman', 8)
                ref_para.space_after = Pt(2)
        else:
            ref_para = doc.add_paragraph()
            ref_run = ref_para.add_run(references)
            set_font(ref_run, 'Times New Roman', 8)

    output_path = os.path.join(OUTPUT_DIR, f"{job_id}_springer.docx")
    doc.save(output_path)
    return output_path


def generate_docx(structured_content: dict, format_type: str, job_id: str) -> str:
    """Main function to generate formatted DOCX documents"""
    format_type = format_type.lower().strip()

    if 'springer' in format_type:
        return generate_springer(structured_content, job_id)
    elif 'ieee' in format_type:
        return generate_ieee(structured_content, job_id)
    elif format_type.startswith('journal:'):
        # Extract journal name e.g. "journal: elsevier" -> "elsevier"
        journal_name = format_type.split(':', 1)[1].strip()
        if 'elsevier' in journal_name:
            return generate_springer(structured_content, job_id)  # Elsevier uses Springer-like single-column
        else:
            # For APA, MLA, Nature, etc. - use IEEE-style academic formatting
            return generate_ieee(structured_content, job_id)
    else:
        # Default to IEEE format
        return generate_ieee(structured_content, job_id)